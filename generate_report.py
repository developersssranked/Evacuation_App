# -*- coding: utf-8 -*-
"""
Генератор отчёта по лабораторной работе "Мобильное приложение Evacuation".
Создаёт файл Отчет.docx в текущей директории.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)


def set_run(run, bold=False, italic=False, size=FONT_SIZE):
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)
    run.font.size = size
    run.bold = bold
    run.italic = italic


def add_paragraph(doc, text, *, bold=False, italic=False, align=None,
                  size=FONT_SIZE, space_after=Pt(6), first_line_indent=Cm(1.25)):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = space_after
    pf.space_before = Pt(0)
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, bold=bold, italic=italic, size=size)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run(r, bold=True, size=Pt(16 if level == 1 else 14))
    return p


def add_centered(doc, text, *, bold=False, italic=False, size=FONT_SIZE,
                 space_after=Pt(6)):
    return add_paragraph(
        doc, text, bold=bold, italic=italic, align=WD_ALIGN_PARAGRAPH.CENTER,
        size=size, space_after=space_after, first_line_indent=None
    )


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(2)
    pf.first_line_indent = None
    r = p.add_run(text)
    set_run(r)
    return p


def add_code(doc, code):
    """Листинг кода — моноширинный шрифт, размер 11."""
    for line in code.split("\n"):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.first_line_indent = None
        pf.left_indent = Cm(0.5)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rFonts.set(qn(attr), "Consolas")
        r.font.size = Pt(11)


def add_screenshot_placeholder(doc, number, description):
    """Рамка-плейсхолдер под скриншот + подпись 'Рисунок N — описание'."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(15)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Рамка вокруг ячейки
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:color"), "808080")
        tcBorders.append(b)
    tcPr.append(tcBorders)
    # Высота "окна"
    for _ in range(6):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
    p_label = cell.paragraphs[3]
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_label.add_run(f"[ Место для скриншота {number} ]")
    set_run(r, italic=True)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    # очистить первый пустой параграф ячейки
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    # Подпись
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.first_line_indent = None
    r = cap.add_run(f"Рисунок {number} — {description}")
    set_run(r, italic=True, size=Pt(13))


def page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)


def add_page_numbers(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run(run, size=Pt(12))
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# =====================================================================

def build():
    doc = Document()
    set_default_style(doc)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    add_page_numbers(doc)

    # ===== ТИТУЛЬНЫЙ ЛИСТ =====
    for _ in range(3):
        add_centered(doc, "")
    add_centered(doc, "МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ", bold=True)
    add_centered(doc, "________________________________________", bold=True)
    add_centered(doc, "(полное наименование учебного заведения)", size=Pt(12))
    add_centered(doc, "")
    add_centered(doc, "Кафедра ________________________________", size=Pt(13))
    for _ in range(6):
        add_centered(doc, "")
    add_centered(doc, "ОТЧЁТ", bold=True, size=Pt(20))
    add_centered(doc, "по лабораторной работе", size=Pt(16))
    add_centered(doc, "")
    add_centered(doc,
        "«Разработка мобильного приложения для построения "
        "маршрута эвакуации к ближайшему укрытию»",
        bold=True, size=Pt(16))
    for _ in range(6):
        add_centered(doc, "")
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(
        "Выполнил(а): студент(ка) гр. _______\n"
        "_____________________ / Ф. И. О. /\n\n"
        "Проверил(а): _______________________\n"
        "_____________________ / Ф. И. О. /"
    )
    set_run(r)
    for _ in range(5):
        add_centered(doc, "")
    add_centered(doc, "Город — 2026", size=Pt(13))
    page_break(doc)

    # ===== СОДЕРЖАНИЕ =====
    add_heading(doc, "СОДЕРЖАНИЕ")
    toc = [
        ("Введение. Цель работы", "3"),
        ("1. Задание", "3"),
        ("2. Теоретическая часть", "4"),
        ("   2.1. Краткие сведения по теме", "4"),
        ("   2.2. Используемые технологии и методы", "4"),
        ("   2.3. Ключевые формулы", "5"),
        ("3. Практическая часть", "6"),
        ("   3.1. Архитектура решения", "6"),
        ("   3.2. Описание модулей и их взаимодействия", "7"),
        ("   3.3. Листинги кода", "8"),
        ("   3.4. Скриншоты интерфейса", "11"),
        ("   3.5. Примеры входных и выходных данных", "13"),
        ("   3.6. Результаты тестирования", "14"),
        ("4. Анализ результатов", "15"),
        ("5. Заключение", "16"),
        ("Список использованных источников", "17"),
    ]
    for title, page in toc:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = None
        pf.space_after = Pt(2)
        tab_stops = pf.tab_stops
        tab_stops.add_tab_stop(Cm(16), WD_ALIGN_PARAGRAPH.RIGHT, leader=2)  # dot leader
        r = p.add_run(f"{title}\t{page}")
        set_run(r)
    page_break(doc)

    # ===== ЦЕЛЬ РАБОТЫ =====
    add_heading(doc, "ЦЕЛЬ РАБОТЫ")
    add_paragraph(doc,
        "Разработать мобильное приложение под платформу Android, "
        "которое в реальном времени отслеживает геопозицию пользователя, "
        "позволяет отмечать на карте точки укрытий и автоматически "
        "строит пешеходный маршрут до ближайшего из них. "
        "Закрепить практические навыки работы с современным стеком "
        "Android (Kotlin, Jetpack Compose, MVVM), картографическим "
        "движком OpenStreetMap (osmdroid) и сервисом маршрутизации OSRM."
    )

    # ===== ЗАДАНИЕ =====
    add_heading(doc, "1. ЗАДАНИЕ")
    add_paragraph(doc, "В рамках лабораторной работы требуется:")
    for t in [
        "Спроектировать архитектуру приложения по шаблону MVVM с реактивным "
        "обменом состоянием между слоями (StateFlow).",
        "Реализовать модуль получения геолокации через Fused Location Provider "
        "и оформить его в виде потока (Flow) обновлений.",
        "Подключить картографический компонент osmdroid и обеспечить "
        "интерактивное отображение карты OpenStreetMap.",
        "Реализовать добавление укрытий на карту по долгому нажатию и из "
        "текущей геопозиции пользователя.",
        "Подключить сервис OSRM для построения пешеходного маршрута и "
        "отобразить маршрут на карте в виде ломаной линии.",
        "Реализовать пользовательский интерфейс на Jetpack Compose: "
        "карточка состояния (GPS, число укрытий, длина и время маршрута), "
        "панель управления, обработка ошибок.",
        "Подготовить инструкцию по запуску приложения и оформить "
        "репозиторий с README.md.",
    ]:
        add_bullet(doc, t)
    page_break(doc)

    # ===== ТЕОРЕТИЧЕСКАЯ ЧАСТЬ =====
    add_heading(doc, "2. ТЕОРЕТИЧЕСКАЯ ЧАСТЬ")

    add_heading(doc, "2.1. Краткие сведения по теме", level=2)
    add_paragraph(doc,
        "Современные мобильные навигационные приложения опираются на три "
        "технологических кита: получение координат устройства от GPS / "
        "сетевых провайдеров, отрисовку карты по векторным или растровым "
        "тайлам, и серверный сервис маршрутизации, который по графу "
        "дорожной сети строит оптимальный путь между двумя точками."
    )
    add_paragraph(doc,
        "Под Android приложение строится на основе компонента Activity, "
        "а вся пользовательская часть в работе реализована декларативно "
        "через Jetpack Compose — современный UI-фреймворк, заменяющий "
        "классический подход с XML-разметкой и императивным управлением "
        "представлениями."
    )
    add_paragraph(doc,
        "Для реактивной связи между слоем данных и слоем представления "
        "применяется паттерн MVVM (Model — View — ViewModel). "
        "ViewModel выступает посредником: подписывается на источники "
        "данных, агрегирует состояние и публикует его через StateFlow, "
        "на которые подписан Compose-UI."
    )

    add_heading(doc, "2.2. Используемые технологии и методы", level=2)
    techs = [
        ("Kotlin 2.x", "основной язык разработки; используются корутины "
                       "и Flow для асинхронных операций."),
        ("Jetpack Compose", "декларативный UI-фреймворк; экраны "
                            "описываются функциями @Composable."),
        ("MVVM + StateFlow", "архитектурный паттерн; состояние UI хранится "
                             "во ViewModel и публикуется как поток."),
        ("Google Play Services — Fused Location Provider",
         "поставщик геолокации, объединяющий данные GPS, Wi-Fi и сотовых "
         "сетей; даёт точность и энергоэффективность лучше прямого GPS."),
        ("osmdroid 6.x", "Android-библиотека для работы с тайлами "
                         "OpenStreetMap; не требует API-ключей."),
        ("OSRM (Open Source Routing Machine)",
         "открытый веб-сервис маршрутизации; используется публичный "
         "endpoint router.project-osrm.org, профиль foot (пешеход)."),
        ("OkHttp 4.x", "HTTP-клиент для синхронных запросов к OSRM."),
        ("Kotlin Coroutines",
         "lifecycle-aware асинхронность; запросы выполняются в "
         "viewModelScope без блокировки UI-потока."),
    ]
    for name, desc in techs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"• {name} — ")
        set_run(r1, bold=True)
        r2 = p.add_run(desc)
        set_run(r2)

    add_heading(doc, "2.3. Ключевые формулы", level=2)
    add_paragraph(doc,
        "Для предварительного выбора ближайшего укрытия (до запроса "
        "к OSRM) используется формула гаверсинусов, дающая расстояние "
        "по дуге большого круга между двумя точками на сфере:"
    )
    add_centered(doc,
        "d = 2 · R · arcsin( √( sin²(Δφ/2) + cos(φ₁)·cos(φ₂)·sin²(Δλ/2) ) )",
        size=Pt(13))
    add_paragraph(doc,
        "где R = 6 371 000 м — радиус Земли, φ — широта, λ — долгота "
        "(в радианах), Δφ и Δλ — разности координат. Формула используется "
        "только для оценки «ближайшей» точки; финальная длина маршрута "
        "берётся из ответа OSRM и учитывает реальные дорожные графы."
    )
    page_break(doc)

    # ===== ПРАКТИЧЕСКАЯ ЧАСТЬ =====
    add_heading(doc, "3. ПРАКТИЧЕСКАЯ ЧАСТЬ")

    add_heading(doc, "3.1. Архитектура решения", level=2)
    add_paragraph(doc,
        "Приложение построено по классической трёхслойной схеме MVVM. "
        "Информационный поток между слоями выглядит следующим образом:"
    )
    add_paragraph(doc,
        "Слой данных (LocationProvider, OsrmClient) → слой "
        "представления модели (EvacuationViewModel, StateFlow) → "
        "слой представления (EvacuationScreen, Composable-функции).",
        italic=True
    )
    add_paragraph(doc,
        "Пользовательские события (нажатие кнопки, долгий тап) приходят "
        "из UI во ViewModel в виде вызовов методов, а обратная связь "
        "(обновлённые координаты, новый маршрут, ошибки) поступает в "
        "UI через подписку Compose на StateFlow."
    )
    add_paragraph(doc,
        "Модули приложения и их назначение приведены в таблице 1."
    )
    # Таблица модулей
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = "Table Grid"
    rows_data = [
        ("Модуль", "Назначение"),
        ("model/Models.kt",
         "Доменные сущности: Coord, Shelter, UserLocation, "
         "функция расстояния haversineMeters."),
        ("location/LocationProvider.kt",
         "Обёртка над Fused Location Provider; "
         "публикует Flow<UserLocation>."),
        ("network/OsrmClient.kt",
         "HTTP-клиент OSRM; возвращает геометрию маршрута, "
         "его длину и время прохождения."),
        ("EvacuationViewModel.kt",
         "Оркестратор: хранит состояние карты, координирует обновления "
         "локации и запросы маршрутов; публикует StateFlow."),
        ("ui/EvacuationScreen.kt",
         "Composable-экран: карта osmdroid, маркеры, маршрут, "
         "карточка статуса, панель управления, диалоги разрешений."),
    ]
    for r_idx, (col1, col2) in enumerate(rows_data):
        cells = tbl.rows[r_idx].cells
        for c, txt in enumerate((col1, col2)):
            cells[c].text = ""
            p = cells[c].paragraphs[0]
            p.paragraph_format.first_line_indent = None
            run = p.add_run(txt)
            set_run(run, bold=(r_idx == 0), size=Pt(13))
    add_centered(doc, "Таблица 1 — Состав модулей приложения",
                 italic=True, size=Pt(13))

    add_heading(doc, "3.2. Описание модулей и их взаимодействия", level=2)
    add_paragraph(doc,
        "LocationProvider оформлен в виде callbackFlow: при подписке "
        "регистрируется LocationCallback Fused-провайдера, при отписке "
        "коллбэк снимается. Каждое обновление координат превращается в "
        "объект UserLocation (широта, долгота, погрешность, время)."
    )
    add_paragraph(doc,
        "OsrmClient выполняет один запрос на построение маршрута: "
        "формирует URL вида /route/v1/foot/lon,lat;lon,lat, разбирает "
        "GeoJSON-геометрию и возвращает список точек маршрута, длину "
        "в метрах и оценочное время в секундах."
    )
    add_paragraph(doc,
        "EvacuationViewModel держит пять StateFlow: список укрытий, "
        "текущая локация пользователя, текущий маршрут, флаг загрузки "
        "маршрута и текст ошибки. Запуск обновлений локации идёт в "
        "viewModelScope; запрос маршрута — также корутина с обработкой "
        "Result для безопасного перехвата сетевых ошибок."
    )
    add_paragraph(doc,
        "EvacuationScreen встраивает MapView через интероп-компонент "
        "AndroidView, подписывается на StateFlow ViewModel и при "
        "изменениях состояния перерисовывает маркеры и линию маршрута. "
        "Запрос runtime-разрешений на геолокацию выполняется через "
        "rememberLauncherForActivityResult."
    )

    add_heading(doc, "3.3. Листинги кода", level=2)
    add_paragraph(doc,
        "Полный исходный код проекта размещён в репозитории GitHub "
        "(ссылка указана в списке источников). Ниже приведены ключевые "
        "фрагменты."
    )

    add_paragraph(doc, "Листинг 1. Доменные модели и формула расстояния "
                       "(model/Models.kt).", italic=True)
    add_code(doc,
"""data class Coord(val lat: Double, val lon: Double)

data class Shelter(
    val id: String = UUID.randomUUID().toString(),
    val name: String,
    val coord: Coord
)

data class UserLocation(
    val coord: Coord,
    val accuracyMeters: Float,
    val timestamp: Long
)

fun haversineMeters(a: Coord, b: Coord): Double {
    val r = 6_371_000.0
    val dLat = Math.toRadians(b.lat - a.lat)
    val dLon = Math.toRadians(b.lon - a.lon)
    val lat1 = Math.toRadians(a.lat)
    val lat2 = Math.toRadians(b.lat)
    val h = sin(dLat / 2).let { it * it } +
            cos(lat1) * cos(lat2) * sin(dLon / 2).let { it * it }
    return 2 * r * atan2(sqrt(h), sqrt(1 - h))
}""")

    add_paragraph(doc, "Листинг 2. Поток обновлений локации "
                       "(location/LocationProvider.kt).", italic=True)
    add_code(doc,
"""@SuppressLint("MissingPermission")
fun updates(intervalMs: Long = 4000L, minIntervalMs: Long = 2000L)
        : Flow<UserLocation> = callbackFlow {
    val client = LocationServices.getFusedLocationProviderClient(context)
    val request = LocationRequest.Builder(
        Priority.PRIORITY_HIGH_ACCURACY, intervalMs
    ).setMinUpdateIntervalMillis(minIntervalMs).build()

    val callback = object : LocationCallback() {
        override fun onLocationResult(result: LocationResult) {
            val loc = result.lastLocation ?: return
            trySend(UserLocation(
                coord = Coord(loc.latitude, loc.longitude),
                accuracyMeters = loc.accuracy,
                timestamp = loc.time
            ))
        }
    }
    client.requestLocationUpdates(request, callback, context.mainLooper)
    awaitClose { client.removeLocationUpdates(callback) }
}""")

    add_paragraph(doc, "Листинг 3. Запрос пешеходного маршрута к OSRM "
                       "(network/OsrmClient.kt).", italic=True)
    add_code(doc,
"""suspend fun route(from: Coord, to: Coord): Result<OsrmRoute> =
    withContext(Dispatchers.IO) {
        runCatching {
            val url = "$baseUrl/route/v1/$profile/" +
                "${'$'}{from.lon},${'$'}{from.lat};" +
                "${'$'}{to.lon},${'$'}{to.lat}" +
                "?overview=full&geometries=geojson"
            val req = Request.Builder().url(url).build()
            client.newCall(req).execute().use { resp ->
                val body = resp.body?.string().orEmpty()
                if (!resp.isSuccessful) error("HTTP ${'$'}{resp.code}")
                val json = JSONObject(body)
                val route = json.getJSONArray("routes").getJSONObject(0)
                val coords = route.getJSONObject("geometry")
                    .getJSONArray("coordinates")
                val points = buildList(coords.length()) {
                    for (i in 0 until coords.length()) {
                        val pair = coords.getJSONArray(i)
                        add(Coord(lat = pair.getDouble(1),
                                  lon = pair.getDouble(0)))
                    }
                }
                OsrmRoute(
                    points = points,
                    distanceMeters = route.getDouble("distance"),
                    durationSeconds = route.getDouble("duration")
                )
            }
        }
    }""")

    add_paragraph(doc, "Листинг 4. Реактивное состояние ViewModel и "
                       "построение маршрута к ближайшему укрытию "
                       "(EvacuationViewModel.kt).", italic=True)
    add_code(doc,
"""private val _shelters = MutableStateFlow<List<Shelter>>(emptyList())
val shelters: StateFlow<List<Shelter>> = _shelters.asStateFlow()

private val _userLocation = MutableStateFlow<UserLocation?>(null)
val userLocation: StateFlow<UserLocation?> = _userLocation.asStateFlow()

private val _route = MutableStateFlow<RouteState?>(null)
val route: StateFlow<RouteState?> = _route.asStateFlow()

fun nearestShelter(): Shelter? {
    val from = _userLocation.value?.coord ?: return null
    return _shelters.value.minByOrNull {
        haversineMeters(from, it.coord)
    }
}

fun buildRouteTo(shelter: Shelter) {
    val from = _userLocation.value?.coord ?: run {
        _error.value = "Нет данных о местоположении"; return
    }
    viewModelScope.launch {
        _isLoadingRoute.value = true
        osrm.route(from, shelter.coord)
            .onSuccess { _route.value = RouteState(
                points = it.points,
                distanceMeters = it.distanceMeters,
                durationSeconds = it.durationSeconds,
                targetShelterId = shelter.id) }
            .onFailure {
                _error.value = "Маршрут не построен: ${'$'}{it.message}"
            }
        _isLoadingRoute.value = false
    }
}""")

    add_paragraph(doc, "Листинг 5. Интеграция MapView в Compose и реакция "
                       "на изменение состояния (ui/EvacuationScreen.kt).",
                  italic=True)
    add_code(doc,
"""@Composable
private fun OsmMap(
    shelters: List<Shelter>,
    userCoord: Coord?,
    routePoints: List<Coord>?,
    onLongPress: (Coord) -> Unit,
    onShelterClick: (Shelter) -> Unit
) {
    val context = LocalContext.current
    val mapView = remember { createMapView(context, onLongPress) }

    DisposableEffect(mapView) {
        mapView.onResume()
        onDispose { mapView.onPause() }
    }

    LaunchedEffect(userCoord) { userCoord?.let {
        updateUserMarker(mapView, it)
    } }
    LaunchedEffect(shelters) {
        updateShelterMarkers(mapView, shelters, onShelterClick)
    }
    LaunchedEffect(routePoints) {
        updateRoutePolyline(mapView, routePoints)
    }

    AndroidView(modifier = Modifier.fillMaxSize(), factory = { mapView })
}""")

    add_heading(doc, "3.4. Скриншоты интерфейса", level=2)
    add_paragraph(doc,
        "На рисунках 1–5 приведены ключевые состояния пользовательского "
        "интерфейса приложения. Все снимки сделаны на эмуляторе Android "
        "Studio (Pixel 6, Android 14) с включённой передачей координат "
        "через Extended Controls → Location."
    )
    add_screenshot_placeholder(doc, 1,
        "запуск приложения и запрос разрешений на доступ к геолокации")
    add_screenshot_placeholder(doc, 2,
        "карта с маркером текущей геопозиции пользователя")
    add_screenshot_placeholder(doc, 3,
        "добавление укрытий на карту долгим нажатием")
    add_screenshot_placeholder(doc, 4,
        "построенный пешеходный маршрут к ближайшему укрытию")
    add_screenshot_placeholder(doc, 5,
        "карточка статуса с координатами, точностью GPS, "
        "числом укрытий и параметрами маршрута")

    add_heading(doc, "3.5. Примеры входных и выходных данных", level=2)
    add_paragraph(doc,
        "Пример входных данных — текущая геопозиция пользователя и "
        "массив координат укрытий, полученный по долгим нажатиям на карту:"
    )
    add_code(doc,
"""Пользователь:
    lat = 55.75582, lon = 37.61730, accuracy = 8.4 м

Укрытия:
    #1  lat = 55.75701, lon = 37.61952
    #2  lat = 55.75412, lon = 37.62110
    #3  lat = 55.75834, lon = 37.61455""")
    add_paragraph(doc,
        "Запрос к OSRM (HTTP GET):"
    )
    add_code(doc,
"""GET https://router.project-osrm.org/route/v1/foot/
    37.61730,55.75582;37.61952,55.75701
    ?overview=full&geometries=geojson""")
    add_paragraph(doc, "Пример ответа OSRM (фрагмент):")
    add_code(doc,
"""{
  "code": "Ok",
  "routes": [{
    "distance": 263.4,
    "duration": 197.5,
    "geometry": {
      "coordinates": [
        [37.61730, 55.75582],
        [37.61774, 55.75618],
        ...
        [37.61952, 55.75701]
      ]
    }
  }]
}""")
    add_paragraph(doc,
        "На UI это отображается как карточка маршрута: «Маршрут: 263 м, "
        "~3 мин», и ломаная линия синего цвета поверх карты."
    )

    add_heading(doc, "3.6. Результаты тестирования", level=2)
    add_paragraph(doc, "Тестирование проводилось на эмуляторе Android "
                       "Studio. Рассмотрены успешные и неудачные сценарии "
                       "(таблица 2).")
    test_rows = [
        ("№", "Сценарий", "Ожидаемый результат", "Итог"),
        ("1", "Запуск приложения, разрешение геолокации",
         "Появляется карта, на ней — синий маркер пользователя",
         "Успешно"),
        ("2", "Длинное нажатие на свободной точке карты",
         "На карте появляется метка укрытия, счётчик увеличивается",
         "Успешно"),
        ("3", "Нажатие на метку укрытия",
         "Строится маршрут от пользователя к укрытию, "
         "выводится длина и время",
         "Успешно"),
        ("4", "Нажатие «К ближайшему» без укрытий",
         "Снэк-бар «Нет добавленных укрытий», маршрут не строится",
         "Успешно"),
        ("5", "Запрос маршрута при выключенном интернете",
         "Снэк-бар «Маршрут не построен: …», UI остаётся откликающимся",
         "Успешно"),
        ("6", "Отказ в предоставлении разрешения",
         "Карта показывается без маркера пользователя; кнопки «Здесь» "
         "и «К ближайшему» заблокированы",
         "Успешно"),
    ]
    tbl2 = doc.add_table(rows=len(test_rows), cols=4)
    tbl2.style = "Table Grid"
    widths = [Cm(1), Cm(5), Cm(7), Cm(2.5)]
    for r_idx, row in enumerate(test_rows):
        cells = tbl2.rows[r_idx].cells
        for c_idx, txt in enumerate(row):
            cells[c_idx].text = ""
            cells[c_idx].width = widths[c_idx]
            p = cells[c_idx].paragraphs[0]
            p.paragraph_format.first_line_indent = None
            run = p.add_run(txt)
            set_run(run, bold=(r_idx == 0), size=Pt(12))
    add_centered(doc, "Таблица 2 — Результаты тестирования",
                 italic=True, size=Pt(13))
    page_break(doc)

    # ===== АНАЛИЗ =====
    add_heading(doc, "4. АНАЛИЗ РЕЗУЛЬТАТОВ")
    add_paragraph(doc,
        "Все поставленные в задании пункты реализованы: "
        "приложение получает геолокацию пользователя в реальном времени, "
        "позволяет добавлять укрытия и строит пешеходный маршрут до "
        "ближайшего из них с помощью внешнего сервиса OSRM. "
        "Реактивность UI обеспечена связкой Compose + StateFlow: "
        "обновление координат сразу сдвигает маркер пользователя, "
        "а добавление укрытия — мгновенно появляется на карте."
    )
    add_paragraph(doc, "В ходе разработки выявлены и решены проблемы:")
    add_bullet(doc,
        "На современных Android (API 29+) тайлы osmdroid не "
        "подгружались из-за scoped storage. Решение — явно указать "
        "osmdroidBasePath и osmdroidTileCache во внутреннем cacheDir "
        "приложения.")
    add_bullet(doc,
        "Стандартный User-Agent osmdroid блокируется тайл-сервером "
        "OSM. Решение — задать значение вида "
        "«Evacuation/1.0 (<packageName>)».")
    add_bullet(doc,
        "Пересоздание MapView при рекомпозиции Compose приводило к "
        "потере состояния. Решение — обернуть создание в remember и "
        "управлять жизненным циклом через DisposableEffect "
        "(onResume / onPause).")
    add_bullet(doc,
        "Запрос маршрута при отсутствии сети ронял приложение. "
        "Решение — обернуть HTTP-вызов в runCatching и проксировать "
        "ошибку в _error через StateFlow.")
    add_paragraph(doc, "Возможные улучшения и оптимизации:")
    add_bullet(doc, "Сохранение списка укрытий в локальной БД "
                    "(Room) и восстановление после перезапуска.")
    add_bullet(doc, "Офлайн-карта: предзагрузка тайлов выбранного "
                    "региона.")
    add_bullet(doc, "Поддержка нескольких профилей маршрутизации "
                    "(пешеход / велосипед / автомобиль).")
    add_bullet(doc, "Push-уведомления при приближении к укрытию "
                    "(geofencing).")
    add_bullet(doc, "Локализация интерфейса (сейчас только русский).")
    page_break(doc)

    # ===== ЗАКЛЮЧЕНИЕ =====
    add_heading(doc, "5. ЗАКЛЮЧЕНИЕ")
    add_paragraph(doc,
        "В результате выполнения лабораторной работы разработано "
        "законченное Android-приложение для построения маршрутов "
        "эвакуации. Все поставленные цели достигнуты: получение "
        "координат через Fused Location Provider, отрисовка карты "
        "OpenStreetMap, маршрутизация через OSRM, реактивный "
        "Compose-интерфейс."
    )
    add_paragraph(doc, "В процессе выполнения работы приобретены навыки:")
    add_bullet(doc, "Проектирование Android-приложений по паттерну "
                    "MVVM с реактивным состоянием.")
    add_bullet(doc, "Использование Jetpack Compose, в том числе "
                    "интеграция с классическими View через AndroidView.")
    add_bullet(doc, "Работа с Kotlin Coroutines и Flow для асинхронных "
                    "операций и потоковых данных.")
    add_bullet(doc, "Подключение и настройка библиотеки osmdroid; "
                    "управление маркерами, полилиниями и тайл-кэшем.")
    add_bullet(doc, "Интеграция с REST-сервисом OSRM, разбор GeoJSON-"
                    "ответов через org.json.")
    add_bullet(doc, "Работа с runtime-разрешениями Android.")
    add_paragraph(doc,
        "В дальнейшем приложение может быть расширено локальным "
        "хранилищем укрытий, офлайн-картами, поддержкой нескольких "
        "транспортных профилей и геофенсингом."
    )
    page_break(doc)

    # ===== ИСТОЧНИКИ =====
    add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    sources = [
        "Репозиторий проекта на GitHub: "
        "https://github.com/<USERNAME>/Evacuation (README.md содержит "
        "инструкции по сборке и запуску).",
        "Официальная документация Android Developers. — URL: "
        "https://developer.android.com (дата обращения: 26.05.2026).",
        "Jetpack Compose — официальное руководство. — URL: "
        "https://developer.android.com/jetpack/compose "
        "(дата обращения: 26.05.2026).",
        "osmdroid: OpenStreetMap-Tools for Android. — URL: "
        "https://github.com/osmdroid/osmdroid "
        "(дата обращения: 26.05.2026).",
        "OSRM API Documentation. Project OSRM. — URL: "
        "https://project-osrm.org/docs/v5.24.0/api/ "
        "(дата обращения: 26.05.2026).",
        "Fused Location Provider API. Google Play Services. — URL: "
        "https://developers.google.com/location-context/fused-location-provider "
        "(дата обращения: 26.05.2026).",
        "Kotlin Coroutines guide. — URL: "
        "https://kotlinlang.org/docs/coroutines-guide.html "
        "(дата обращения: 26.05.2026).",
    ]
    for i, s in enumerate(sources, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{i}. {s}")
        set_run(r, size=Pt(13))

    out = "Отчет.docx"
    doc.save(out)
    print(f"OK: создан {out}")


if __name__ == "__main__":
    build()
